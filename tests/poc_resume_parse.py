"""
POC: Resume Parsing Core — proves the entire parsing pipeline in isolation.
1. Generates sample resumes (2 PDFs varied formats, 1 DOCX) as test fixtures
2. Extracts text (pdfplumber for PDF, python-docx for DOCX)
3. Sends text to OpenAI (Emergent LLM key) with strict JSON schema prompt
4. Validates structured output + confidence flags
5. Tests bulk parsing (all files in one run, concurrent)
"""
import asyncio
import json
import os
import re
import sys
from pathlib import Path

from dotenv import load_dotenv

load_dotenv("/app/backend/.env")

FIXTURES = Path("/app/tests/fixtures")
FIXTURES.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# 1. Generate sample resumes
# ---------------------------------------------------------------------------

def make_pdf_resume_1():
    """Standard, well-structured resume PDF."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
    p = FIXTURES / "resume_sarah_chen.pdf"
    c = canvas.Canvas(str(p), pagesize=LETTER)
    y = 750
    lines = [
        ("Helvetica-Bold", 16, "Sarah Chen"),
        ("Helvetica", 10, "sarah.chen@example.com | (415) 555-0192 | San Francisco, CA"),
        ("Helvetica", 10, "Senior Software Engineer at CloudScale Inc."),
        ("Helvetica-Bold", 12, "EXPERIENCE"),
        ("Helvetica-Bold", 10, "Senior Software Engineer — CloudScale Inc."),
        ("Helvetica", 9, "March 2021 - Present | San Francisco, CA"),
        ("Helvetica", 9, "Led backend team building distributed data pipelines in Python and Go."),
        ("Helvetica-Bold", 10, "Software Engineer — DataWorks LLC"),
        ("Helvetica", 9, "June 2017 - February 2021 | Seattle, WA"),
        ("Helvetica", 9, "Built REST APIs with FastAPI and PostgreSQL; containerized services with Docker."),
        ("Helvetica-Bold", 12, "EDUCATION"),
        ("Helvetica", 10, "B.S. Computer Science — University of Washington, 2013 - 2017"),
        ("Helvetica-Bold", 12, "SKILLS"),
        ("Helvetica", 9, "Python, Go, FastAPI, PostgreSQL, Docker, Kubernetes, AWS, React"),
    ]
    for font, size, text in lines:
        c.setFont(font, size)
        c.drawString(72, y, text)
        y -= 22
    c.save()
    return p


def make_pdf_resume_2():
    """Messier resume PDF: missing phone, unusual layout."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
    p = FIXTURES / "resume_miguel_torres.pdf"
    c = canvas.Canvas(str(p), pagesize=LETTER)
    y = 750
    lines = [
        ("Helvetica-Bold", 14, "MIGUEL TORRES  //  Product Designer"),
        ("Helvetica", 9, "miguel.torres.design@example.io"),
        ("Helvetica", 9, "Portfolio: miguelt.example.io  ::  Austin"),
        ("Helvetica-Bold", 11, "Work"),
        ("Helvetica", 9, "Brightpath Studio -- Lead Product Designer (2022-now)"),
        ("Helvetica", 9, "Owned design system; shipped mobile app used by 2M users."),
        ("Helvetica", 9, "Freelance -- UX Designer (2019-2022)"),
        ("Helvetica", 9, "Clients incl. fintech + healthtech startups."),
        ("Helvetica-Bold", 11, "Learning"),
        ("Helvetica", 9, "BFA Graphic Design, RISD (2015-2019)"),
        ("Helvetica-Bold", 11, "Tools"),
        ("Helvetica", 9, "Figma / Sketch / Framer / HTML-CSS / prototyping / user research"),
    ]
    for font, size, text in lines:
        c.setFont(font, size)
        c.drawString(72, y, text)
        y -= 20
    c.save()
    return p


def make_docx_resume():
    """DOCX resume."""
    import docx
    p = FIXTURES / "resume_priya_patel.docx"
    d = docx.Document()
    d.add_heading("Priya Patel", 0)
    d.add_paragraph("Email: priya.patel@example.org | Phone: +1-312-555-0147 | Chicago, IL")
    d.add_paragraph("Engineering Manager, FinEdge Technologies")
    d.add_heading("Professional Experience", level=1)
    d.add_paragraph("Engineering Manager — FinEdge Technologies (Jan 2020 – Present)")
    d.add_paragraph("Manage a team of 12 engineers across payments and risk platforms.")
    d.add_paragraph("Senior Developer — BankSoft Corp (Aug 2015 – Dec 2019)")
    d.add_paragraph("Developed core banking microservices in Java and Spring Boot.")
    d.add_heading("Education", level=1)
    d.add_paragraph("M.S. Software Engineering — Illinois Institute of Technology (2013 – 2015)")
    d.add_paragraph("B.Tech Computer Engineering — University of Mumbai (2009 – 2013)")
    d.add_heading("Skills", level=1)
    d.add_paragraph("Java, Spring Boot, Microservices, Team Leadership, Agile, SQL, Kafka")
    d.save(str(p))
    return p


# ---------------------------------------------------------------------------
# 2. Text extraction
# ---------------------------------------------------------------------------

def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    text = ""
    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = [pg.extract_text() or "" for pg in pdf.pages]
        text = "\n".join(pages)
    elif ext == ".docx":
        import docx
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# 3. LLM structured extraction
# ---------------------------------------------------------------------------

PARSE_SYSTEM_PROMPT = """You are an expert resume parser. You receive raw text extracted from a resume (PDF/DOCX) and return ONLY a valid JSON object — no markdown fences, no commentary.

JSON schema:
{
  "name": string|null,
  "email": string|null,
  "phone": string|null,
  "current_title": string|null,
  "current_company": string|null,
  "location": string|null,
  "experience": [{"company": string, "title": string, "start_date": string|null, "end_date": string|null, "description": string|null}],
  "education": [{"school": string, "degree": string|null, "start_date": string|null, "end_date": string|null}],
  "skills": [string],
  "confidence": {"name": number, "email": number, "phone": number, "current_title": number, "current_company": number, "location": number}
}

Rules:
- If a field cannot be confidently determined, set it to null (or [] for arrays). NEVER guess or fabricate.
- confidence values are 0.0-1.0 reflecting how certain you are of each extracted top-level field. Use 0.0 for null fields.
- Dates as human-readable strings, e.g. "Mar 2021", "2019", "Present".
- Skills: individual skill strings, deduplicated.
- Output raw JSON only."""


async def parse_resume_text(text: str, file_label: str) -> dict:
    from emergentintegrations.llm.chat import LlmChat, UserMessage

    api_key = os.environ["EMERGENT_LLM_KEY"]
    chat = LlmChat(
        api_key=api_key,
        session_id=f"poc-parse-{file_label}",
        system_message=PARSE_SYSTEM_PROMPT,
    ).with_model("openai", "gpt-5.4-mini")

    msg = UserMessage(text=f"Parse this resume text:\n\n{text[:15000]}")
    resp = await chat.send_message(msg)
    raw = resp if isinstance(resp, str) else str(resp)
    # Strip potential markdown fences
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Retry once with stricter instruction
        msg2 = UserMessage(text="Your previous output was not valid JSON. Return ONLY the raw JSON object, nothing else.")
        resp2 = await chat.send_message(msg2)
        raw2 = (resp2 if isinstance(resp2, str) else str(resp2)).strip()
        if raw2.startswith("```"):
            raw2 = re.sub(r"^```(?:json)?\s*", "", raw2)
            raw2 = re.sub(r"\s*```$", "", raw2)
        return json.loads(raw2)


def low_confidence_fields(parsed: dict, threshold: float = 0.7) -> list:
    conf = parsed.get("confidence", {}) or {}
    return [f for f, v in conf.items() if (v is None or v < threshold)]


# ---------------------------------------------------------------------------
# 4. Test runner
# ---------------------------------------------------------------------------

REQUIRED_KEYS = {"name", "email", "phone", "current_title", "current_company",
                 "location", "experience", "education", "skills", "confidence"}

EXPECTATIONS = {
    "resume_sarah_chen.pdf": {"name": "Sarah Chen", "email": "sarah.chen@example.com", "min_exp": 2, "min_edu": 1, "min_skills": 5},
    "resume_miguel_torres.pdf": {"name": "Miguel Torres", "email": "miguel.torres.design@example.io", "min_exp": 2, "min_edu": 1, "min_skills": 3, "phone_should_be_null": True},
    "resume_priya_patel.docx": {"name": "Priya Patel", "email": "priya.patel@example.org", "min_exp": 2, "min_edu": 2, "min_skills": 5},
}


async def run_one(path: Path):
    label = path.name
    text = extract_text(str(path))
    assert len(text) > 100, f"{label}: extracted text too short ({len(text)} chars)"
    parsed = await parse_resume_text(text, label)
    missing = REQUIRED_KEYS - set(parsed.keys())
    assert not missing, f"{label}: missing keys {missing}"

    exp = EXPECTATIONS[label]
    errors = []
    if exp["name"].lower() not in (parsed.get("name") or "").lower():
        errors.append(f"name mismatch: {parsed.get('name')}")
    if (parsed.get("email") or "").lower() != exp["email"]:
        errors.append(f"email mismatch: {parsed.get('email')}")
    if len(parsed.get("experience") or []) < exp["min_exp"]:
        errors.append(f"experience too few: {len(parsed.get('experience') or [])}")
    if len(parsed.get("education") or []) < exp["min_edu"]:
        errors.append(f"education too few: {len(parsed.get('education') or [])}")
    if len(parsed.get("skills") or []) < exp["min_skills"]:
        errors.append(f"skills too few: {len(parsed.get('skills') or [])}")
    if exp.get("phone_should_be_null") and parsed.get("phone"):
        errors.append(f"phone should be null/absent but got: {parsed.get('phone')}")

    lc = low_confidence_fields(parsed)
    return label, parsed, errors, lc


async def main():
    print("=" * 70)
    print("POC: Resume Parsing Core")
    print("=" * 70)

    files = [make_pdf_resume_1(), make_pdf_resume_2(), make_docx_resume()]
    print(f"[1] Generated {len(files)} fixture resumes: {[f.name for f in files]}")

    # Text extraction check
    for f in files:
        t = extract_text(str(f))
        print(f"[2] Extracted {len(t)} chars from {f.name}")

    # Bulk parse (concurrent) — simulates bulk upload
    print("[3] Parsing all resumes concurrently via LLM...")
    results = await asyncio.gather(*[run_one(f) for f in files], return_exceptions=True)

    all_pass = True
    for r in results:
        if isinstance(r, Exception):
            all_pass = False
            print(f"  FAIL (exception): {r!r}")
            continue
        label, parsed, errors, lc = r
        status = "PASS" if not errors else "FAIL"
        if errors:
            all_pass = False
        print(f"\n  [{status}] {label}")
        print(f"    name={parsed.get('name')} | email={parsed.get('email')} | phone={parsed.get('phone')}")
        print(f"    title={parsed.get('current_title')} @ {parsed.get('current_company')} | loc={parsed.get('location')}")
        print(f"    exp={len(parsed.get('experience') or [])} | edu={len(parsed.get('education') or [])} | skills={len(parsed.get('skills') or [])}")
        print(f"    low_confidence_fields={lc}")
        for e in errors:
            print(f"    ERROR: {e}")

    print("\n" + "=" * 70)
    print("RESULT:", "ALL TESTS PASSED ✔" if all_pass else "SOME TESTS FAILED ✘")
    print("=" * 70)
    sys.exit(0 if all_pass else 1)


if __name__ == "__main__":
    asyncio.run(main())
