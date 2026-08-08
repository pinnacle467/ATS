"""Resume parsing core — proven in POC (/app/tests/poc_resume_parse.py)."""
import io
import json
import re

from grok_client import grok_json
from industry_taxonomy import INDUSTRY_PROMPT_RULES, normalize_industry_list

PARSE_SYSTEM_PROMPT = f"""You are an expert resume parser. You receive raw text extracted from a resume (PDF/DOCX) and return ONLY a valid JSON object — no markdown fences, no commentary.

JSON schema:
{{
  "name": string|null,
  "email": string|null,
  "phone": string|null,
  "current_title": string|null,
  "current_company": string|null,
  "location": string|null,
  "notice_period": string|null,
  "experience": [{{"company": string, "title": string, "start_date": string|null, "end_date": string|null, "description": string|null}}],
  "education": [{{"school": string, "degree": string|null, "start_date": string|null, "end_date": string|null}}],
  "skills": [string],
  "industry": [string],
  "confidence": {{"name": number, "email": number, "phone": number, "current_title": number, "current_company": number, "location": number}}
}}

Rules:
- If a field cannot be confidently determined, set it to null (or [] for arrays). NEVER guess or fabricate.
- confidence values are 0.0-1.0 reflecting how certain you are of each extracted top-level field. Use 0.0 for null fields.
- notice_period: only fill if the resume explicitly states an availability/notice period (e.g. "Notice period: 30 days", "Available immediately"). Otherwise null — do not guess.
- Dates as human-readable strings, e.g. "Mar 2021", "2019", "Present".
- Skills: individual skill strings, deduplicated.
- {INDUSTRY_PROMPT_RULES}
- Output raw JSON only."""


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    text = ''
    if ext == 'pdf':
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [pg.extract_text() or '' for pg in pdf.pages]
        text = '\n'.join(pages)
    elif ext in ('docx',):
        import docx
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append(' | '.join(cell.text for cell in row.cells))
        text = '\n'.join(parts)
    elif ext == 'txt':
        text = data.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f'Unsupported file type: .{ext} (use PDF or DOCX)')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _strip_fences(raw: str) -> str:
    raw = raw.strip()
    if raw.startswith('```'):
        raw = re.sub(r'^```(?:json)?\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)
    return raw


EMPTY_PARSE = {
    'name': None, 'email': None, 'phone': None, 'current_title': None,
    'current_company': None, 'location': None, 'notice_period': None, 'experience': [], 'education': [],
    'skills': [], 'industry': [], 'confidence': {},
}


async def parse_resume_text(text: str, session_label: str) -> dict:
    try:
        parsed = await grok_json(
            system=PARSE_SYSTEM_PROMPT,
            user=f'Parse this resume text:\n\n{text[:15000]}',
            reasoning_effort='none',
            max_tokens=6000,
        )
    except json.JSONDecodeError:
        return {**EMPTY_PARSE}
    merged = {**EMPTY_PARSE, **{k: v for k, v in parsed.items() if k in EMPTY_PARSE}}
    merged['industry'] = normalize_industry_list(merged.get('industry'))
    return merged


def low_confidence_fields(parsed: dict, threshold: float = 0.7):
    conf = parsed.get('confidence') or {}
    return [f for f, v in conf.items() if (v is None or v < threshold)]
